"""
Generates docs/aic_v2_guide.docx — the AIC v2 walkthrough + technical
+ style guide, mirroring the data_ingester three-chapter format.

Run from the project root:
    /path/to/python docs/generate_guide.py

Re-running overwrites the existing .docx.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

OUT_PATH = Path(__file__).parent / "aic_v2_guide.docx"


# ── Document setup ───────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


# ── Helpers ──────────────────────────────────────────────────────────────────

def h1(text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x4E, 0x10, 0x6F)


def h2(text: str) -> None:
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x3A, 0x0D, 0x54)


def h3(text: str) -> None:
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x08, 0x40)


def p(text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(6)


def bullet(text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def numbered(text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def code(text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.15)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(8)

    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F4F2F8")
    pPr.append(shd)

    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2A, 0x1A, 0x3A)


def callout(text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.2)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(8)

    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),  "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "FFF8E1")
    pPr.append(shd)

    run = para.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x40, 0x10)


def page_break() -> None:
    from docx.enum.text import WD_BREAK
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


# ── COVER ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("AIC v2 — Engineering Walkthrough")
tr.font.size = Pt(28)
tr.bold = True
tr.font.color.rgb = RGBColor(0x4E, 0x10, 0x6F)

sub = doc.add_paragraph()
sr = sub.add_run("FastAPI + Next.js replacement for the AIC_PHASE1 Streamlit UI · "
                 "User guide · Technical reference · Style guide")
sr.font.size = Pt(13)
sr.italic = True
sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p("")

callout(
    "This guide mirrors the data_ingester three-chapter format. Chapter 1 is for "
    "analysts running the tool. Chapter 2 is for developers picking up the codebase. "
    "Chapter 3 is the design-system reference for adding new pages without breaking "
    "visual consistency."
)

page_break()


# ═════════════════════════════════════════════════════════════════════════════
#                              CHAPTER 1
# ═════════════════════════════════════════════════════════════════════════════

h1("Chapter 1 — User Guide & Walkthrough")

# ── What this tool does ──────────────────────────────────────────────────────
h2("What This Tool Does")

p(
    "AIC (Assortment Intelligence Classifier) maps retail product attributes "
    "from raw flat-file inputs to clean, labelled outputs. The pipeline runs in "
    "three stages: a deterministic lookup against historical data, a BM25 text "
    "match for novel descriptions, and an XGBoost classifier as a learned "
    "fallback. An ensemble vote produces the final predictions, which the analyst "
    "reviews in a QC wizard before downloading the cleaned workbook."
)

p(
    "Phase 2 / 3 picks up where Phase 1 finishes — it takes the QC'd workbook "
    "plus a project metadata bundle, runs the downstream processing pipeline, "
    "pauses for any BRAND vs TOOL_BRAND mismatches that need analyst review, "
    "and outputs a category-split export."
)

p(
    "AIC v2 is the FastAPI + Next.js replacement for the original Streamlit "
    "wrapper. The ML code is identical (the ml_package and phase3_package are "
    "vendored unchanged from upstream); only the UI and the run-management "
    "layer have been rebuilt for stability under concurrent use, better error "
    "messaging, and a more modern feel."
)

callout(
    "If you already used the Streamlit version: every feature is preserved. "
    "What's new is queue position + ETA when the pipeline is busy, deep-linkable "
    "run URLs, friendly remediation dialogs instead of stack traces, and a "
    "drag-and-drop file upload."
)

# ── Launching the app ────────────────────────────────────────────────────────
h2("Launching the App")

h3("Option 1 — Docker Desktop (recommended for testing)")

p("The repository ships with a Dockerfile. Build and run:")
code(
    "cd ~/Downloads/AIC\n"
    "docker build -t aic_v2:latest .\n"
    "docker run -p 8000:8000 aic_v2:latest"
)
p("Then open http://localhost:8000 in your browser. The image is fully "
  "self-contained — Python deps, NLTK corpora, and the static frontend all ship "
  "inside it. No internet is required at runtime, so the same image runs on a "
  "walled-garden server without modification.")

h3("Option 2 — Local dev mode (for editing code)")

p("Two terminals, two processes:")
code(
    "# Terminal 1 — FastAPI BFF\n"
    "cd ~/Downloads/AIC\n"
    "python -m venv .venv && source .venv/bin/activate\n"
    "pip install -r requirements.txt\n"
    "uvicorn api.main:app --port 8000 --reload\n"
    "\n"
    "# Terminal 2 — Next.js dev server\n"
    "cd web\n"
    "npm install\n"
    "NEXT_PUBLIC_API_URL=http://localhost:8000 npm run dev"
)
p("Open http://localhost:3000. Code edits hot-reload in both processes.")

# ── Walkthrough: Phase 1 ─────────────────────────────────────────────────────
h2("Phase 1 — Attribute Mapping Walkthrough")

p(
    "We'll use the example files shipped in examples/ to walk through the full "
    "Phase 1 → QC → Finalise → Phase 2 flow. The data is small (60 history rows, "
    "8 new products) so each step takes seconds, but every match path is "
    "exercised so you'll see realistic predictions."
)

h3("Step 1 — Upload the inputs")

p(
    "On the Phase 1 page, the default mode is 'Individual files'. Click 'Choose…' "
    "(or drag-and-drop) and pick:"
)
bullet("examples/AIC_Phase1_input.xlsx — META + FINAL sheets")
bullet("examples/AIC_Phase1_flat_file.csv — the new products to classify")

p(
    "The drop target highlights brand-purple while you drag. Once both files are "
    "selected the green 'Run pipeline' button activates."
)

callout(
    "Alternative: switch to 'Project ZIP' mode and upload a single zip "
    "containing the xlsx + csv + (optionally) the three Phase 2 .txt files. "
    "After Phase 1 finishes, the same project files can feed Phase 2 directly "
    "without re-uploading."
)

h3("Step 2 — Watch the pipeline run")

p(
    "Click 'Run pipeline'. The page replaces the upload form with a progress "
    "panel. You'll see four stages tick through:"
)
numbered("Lookup — deterministic match against historical data")
numbered("Text match — BM25 for novel descriptions")
numbered("ML — XGBoost classifier")
numbered("Ensemble — combines all three signals")

p(
    "The log box below the progress bar streams every line the pipeline emits, "
    "colour-coded by severity. If any error or warning surfaces, a chip at the "
    "top counts them."
)

callout(
    "If another analyst is already running a pipeline, your run waits in a queue. "
    "You'll see a chip that says 'Queued — position 2 of 3' along with an ETA "
    "computed from the median of recent runs. The pipeline serialises one run "
    "at a time per server (this is a stability decision — see Chapter 2)."
)

h3("Step 3 — Review predictions in the QC wizard")

p(
    "When the pipeline finishes Phase 1 you're redirected to the QC review "
    "wizard at /qc?runId=…. Two sheets appear in the example data: BRAND and "
    "PACK_SIZE."
)

p(
    "For each sheet you see a grid where every new product has its prediction. "
    "Cells are colour-coded:"
)
bullet("White / pale — high-confidence prediction, leave alone")
bullet("Lavender — you've edited this row")
bullet("Pink — high-priority review (low score, no ML agreement)")
bullet("Yellow — low score with no ML model agreement")

p(
    "Click any prediction cell to drop down to the list of valid attribute "
    "values (sourced from META). Pick a different one if the prediction is "
    "wrong. Edits auto-save with a 600 ms debounce — you don't need a 'save' "
    "button between cells."
)

h3("Step 4 — Save & Finalise")

p(
    "When you're happy with one sheet, click 'Save & Next' to advance. On the "
    "last sheet the button changes to 'Save & Finalize' — clicking that writes "
    "the final workbook and shows the download card."
)

h3("Step 5 — Download or hand off")

p(
    "The Done screen offers four actions:"
)
bullet("Download File_For_Mapping_QC.xlsx — the final QC'd workbook")
bullet("Continue to Phase 2 / 3 → — skip the re-upload, hand the in-memory artifact directly to Phase 2")
bullet("Download bundle — zip with workbook + log + your edits as JSON + run metadata")
bullet("Download log — plain-text log of every line the pipeline emitted")

callout(
    "The bundle is the audit trail. If you're handing off the output to "
    "downstream consumers, the bundle is what proves which decisions you made."
)

# ── Walkthrough: Phase 2 ─────────────────────────────────────────────────────
h2("Phase 2 / 3 — Pipeline & QC Walkthrough")

p(
    "Phase 2 takes the QC'd workbook plus three project metadata files "
    "(ModelInfo.txt, Attributes.txt, AttributeValues.txt), runs the downstream "
    "phase3_package pipeline, and outputs a Cleaned Output workbook with a "
    "post-QC re-collapse step that splits the result into per-category CSVs."
)

h3("Three ways to start a Phase 2 run")

bullet("From a Phase 1 run — clicking 'Continue to Phase 2 / 3 →' on the Phase 1 Done screen passes the artifacts directly. No re-upload.")
bullet("Project zip — upload a single zip containing all four files.")
bullet("Individual files — upload xlsx + the three .txt files separately.")

h3("Configuration")

p(
    "Below the upload area is a configuration row with three controls:"
)
bullet("RAW UPC10 column — auto-populated from the workbook's columns. The dropdown shows the columns that look like UPC fields.")
bullet("Custom SKU collapse — toggle to use the project-specific collapse logic.")
bullet("Skip RMRR tagging — toggle if the input doesn't include RAW_US_MULTI_RETAILER_RESTRICTED.")

p(
    "Below that is the Advanced configuration disclosure with two sections:"
)
bullet("Private label rules — per-retailer (walmart / cvs / heb) enable + label, defaults match the production pipeline.")
bullet("Brand override rules — always shown, with one ready-to-fill row by default. Force-map specific (manufacturer, brand) pairs to a different TOOL_BRAND. Each cell is a dropdown sourced from the autodetect, with a 'custom' fallback to free-text values.")

h3("Mismatch review")

p(
    "If the pipeline detects rows where BRAND and TOOL_BRAND disagree, the run "
    "pauses at state 'mismatch_pending' and shows a wizard. Each group is one "
    "model run's worth of mismatches. For each mismatched (BRAND, TOOL_BRAND) "
    "pair you pick the corrected value from a dropdown, then advance to the "
    "next group. Rows that match an expected pattern (PRIVATE LABEL / "
    "RESTRICTED / EXCLUDE) are greyed and can be left alone."
)

p(
    "Submitting the corrections wakes the worker, which applies them and runs "
    "Phase B to completion."
)

h3("Output and post-QC")

p(
    "When the run finishes you can:"
)
bullet("Download output.xlsx — the Cleaned Output workbook")
bullet("Edit it externally and re-upload via 'Post-QC: edit & re-upload' to trigger a re-collapse + per-category CSV split")
bullet("Download the resulting AIC_Phase2_3_exports.zip")
bullet("Download the run bundle — same audit trail as Phase 1, plus mismatch corrections")

# ── Understanding the data files ─────────────────────────────────────────────
h2("Understanding the Data Files")

h3("Phase 1 input — Excel (META + FINAL)")

p(
    "META is a configuration sheet that tells the pipeline which raw input columns "
    "key off which output attribute. Each row has:"
)
bullet("Attribute Name in MDM — the raw column name in your CSV / FINAL sheet (e.g. ITEM_DESC, BRAND_RAW)")
bullet("Attribute Group name — the output attribute it feeds into (e.g. BRAND, PACK_SIZE)")
bullet("Attribute_Type — set to MODELING for attributes the pipeline should classify")
bullet("Type — usually blank; reserved for downstream type hints")

p(
    "FINAL is the historical training data. Each row is one previously-labelled "
    "product, with the raw key columns and the correct output labels. The "
    "lookup stage aggregates over identical key combinations; XGBoost trains on "
    "this corpus."
)

h3("Phase 1 input — CSV flat file")

p(
    "One row per new product to classify. Must contain the same raw key columns "
    "the META sheet references, plus enough Phase-2 metadata that a chained "
    "Phase 2 run can run end-to-end (UPC10, SKU, RAW_BRAND, RAW_MANUFACTURER, "
    "ASSORTMENT_CATEGORY_DEFINITION, etc.)."
)

h3("Phase 1 output — File_For_Mapping_QC.xlsx")

p(
    "The workbook the QC wizard finalises has these sheets:"
)
bullet("FINAL — the original training data (carried through unchanged)")
bullet("FLAT_FILE_OUT — the new-product input plus predicted attribute columns")
bullet("META — the original META sheet (carried through)")
bullet("Final_BRAND_lkp, Final_PACK_SIZE_lkp, … — one per attribute, capturing every prediction along with its lookup, BM25, and XGB scores. This is what the QC wizard surfaces and edits.")

h3("Phase 2 input — three .txt files")

p(
    "phase3_package needs a project-specific metadata triplet alongside the QC "
    "workbook:"
)
bullet("ModelInfo.txt — pipe-delimited, one row per model with its category and configuration")
bullet("Attributes.txt — defines each attribute's id, type, and metadata")
bullet("AttributeValues.txt — the canonical list of valid values per attribute")

# ── When things go wrong ─────────────────────────────────────────────────────
h2("When Things Go Wrong")

p(
    "If a run fails, the page replaces the progress panel with a remediation "
    "dialog. The dialog is colour-coded by category:"
)
bullet("Amber — input issue (missing sheet, missing project file). The advice tells you which file is missing or which sheet to add.")
bullet("Amber — config issue (column name doesn't match the workbook). The advice tells you to check the advanced config.")
bullet("Red — server error. The advice tells you to download the log and contact support.")

p(
    "Every dialog has a 'Show technical detail' toggle that exposes the full "
    "traceback for support, plus a 'Download log' link and a 'Start over' "
    "button. The technical detail is preserved alongside the user-facing "
    "advice — analysts see the friendly version, support sees the trace."
)

# ── Maintainability ──────────────────────────────────────────────────────────
h2("Maintainability — Refreshing from Upstream")

p(
    "The ml_package, phase3_package, and aic_utils.py modules are vendored "
    "verbatim from github.com/Ciaran-MacDermott/AIC_PHASE1@main. When upstream "
    "changes (new model, new transformation, etc.):"
)
numbered("Replace the three vendored items with the new versions")
numbered("Run the integration test suite — it exercises the real pipeline end-to-end against synthetic fixtures")
numbered("If anything new is added that hits NLTK at import, bundle the corpus in nltk_data/ and confirm the bootstrap test still passes")
numbered("Build the Docker image and smoke-test against examples/")
numbered("Commit + push")

callout(
    "The api/ layer is the only thing you should edit in this repo. Treat ml_package "
    "and phase3_package as read-only. Any pipeline behaviour change goes upstream "
    "to AIC_PHASE1 first."
)

page_break()


# ═════════════════════════════════════════════════════════════════════════════
#                              CHAPTER 2
# ═════════════════════════════════════════════════════════════════════════════

h1("Chapter 2 — Technical Guide for Developers")

# ── Project overview ─────────────────────────────────────────────────────────
h2("1. Project Overview and Philosophy")

p(
    "AIC v2 replaces the Streamlit wrapper in github.com/Ciaran-MacDermott/AIC_PHASE1 "
    "with a FastAPI BFF (Backend For Frontend) plus a Next.js static-export "
    "frontend. The ML code is unchanged — the entire ml_package, phase3_package, "
    "and aic_utils.py are vendored from upstream and the contract is: don't "
    "modify the vendored code. Any pipeline change goes upstream first."
)

p("Three architectural decisions drive the design:")

bullet(
    "Runs are background jobs, not synchronous POSTs. Pipelines take minutes. "
    "Worker thread + in-memory JobRegistry + status polling. One process-wide "
    "PIPELINE_LOCK so concurrent runs don't trample each other through the "
    "legacy code's chdir / sys.path mutations."
)
bullet(
    "Phase 1 and Phase 2 are separate JobRecords linked by parent_run_id. "
    "Their state machines differ (only Phase 2 has mismatch_pending), inputs "
    "differ, terminal artifacts differ. One unified record would force shared "
    "semantics that don't fit either phase cleanly."
)
bullet(
    "QC sheet payloads are self-contained. The server pre-computes row_flags "
    "(priority / low-score-no-ML) and ships original_values so the React grid "
    "does cellStyle + edit detection client-side without per-keystroke "
    "round-trips. Saves are diff-only ({row_id, attribute_value}) and held "
    "in memory until finalise writes the xlsx via ml_package.write_results."
)

# ── Python environment setup ─────────────────────────────────────────────────
h2("2. Python Environment Setup")

p("Python 3.10+ is required (xgboost wheels and modern type hints).")

code(
    "# Create a venv\n"
    "python -m venv .venv\n"
    "source .venv/bin/activate    # macOS / Linux\n"
    ".venv\\Scripts\\activate     # Windows\n"
    "\n"
    "# Install runtime + test deps\n"
    "pip install -r requirements.txt"
)

p(
    "Heavy ML deps (xgboost, scikit-learn, pandas, nltk) take a couple minutes "
    "on first install. The .venv is gitignored — never commit it. Tests stub "
    "ml_package by default so the fast suite (~80 tests) doesn't need the heavy "
    "stack at import time."
)

# ── FastAPI fundamentals ─────────────────────────────────────────────────────
h2("3. FastAPI Fundamentals")

p(
    "FastAPI is an async-capable Python web framework with automatic OpenAPI "
    "schemas, Pydantic validation, and dependency injection. We use a small "
    "subset:"
)

code(
    "from fastapi import FastAPI, UploadFile, HTTPException\n"
    "from pydantic import BaseModel\n"
    "\n"
    "app = FastAPI(title='AIC API')\n"
    "\n"
    "class RunCreated(BaseModel):\n"
    "    run_id: str\n"
    "\n"
    "@app.post('/api/phase1/runs', response_model=RunCreated)\n"
    "async def create_phase1_run(xlsx: UploadFile, csv: UploadFile) -> RunCreated:\n"
    "    if not xlsx.filename.endswith('.xlsx'):\n"
    "        raise HTTPException(400, 'xlsx file must be .xlsx')\n"
    "    record = jobs.registry.create(phase='phase1', tmpdir=...)\n"
    "    worker.start_phase1(record, ...)\n"
    "    return RunCreated(run_id=record.run_id)"
)

p(
    "Pydantic models double as request validation, response serialisation, and "
    "the source of truth for the TypeScript types the frontend consumes "
    "(manually mirrored in web/lib/types.ts — codegen is a v2 nice-to-have)."
)

# ── Background workers ───────────────────────────────────────────────────────
h2("4. Background Workers, Locks, and Pause/Resume")

p(
    "When a Phase 1 or Phase 2 run starts, FastAPI hands the request off to a "
    "background thread and returns immediately with the run_id. The worker "
    "thread does the actual ML work. The HTTP layer reads the worker's "
    "JobRecord through getters that snapshot under a per-record lock, so "
    "concurrent polls never see torn state."
)

h3("PIPELINE_LOCK")

p(
    "The legacy ml_package code mutates global state — it changes the working "
    "directory mid-run and adds entries to sys.path. Two pipelines running at "
    "once would race on these mutations. The fix is a single process-wide lock:"
)

code(
    "PIPELINE_LOCK = threading.Lock()\n"
    "\n"
    "def run_phase1_worker(record, excel_path, csv_path):\n"
    "    set_state(record, state='running', stage_label='Waiting for pipeline lock…')\n"
    "    with PIPELINE_LOCK:\n"
    "        set_state(record, stage_label='Starting…')\n"
    "        # … run the pipeline\n"
    "        set_state(record, state='qc_ready')"
)

p(
    "While a run holds the lock, others queue up at state='running' but stage_label="
    "'Waiting for pipeline lock…'. The frontend surfaces this as a queue chip "
    "with position and ETA — see Chapter 1 for the user-facing UX."
)

callout(
    "Long-term the right fix is subprocess isolation per run — that's the v2 "
    "scaling path. For now, a single shared process is fine for ~20 concurrent "
    "users with clear queue feedback."
)

h3("Pause/resume via threading.Event")

p(
    "Phase 2 hits a pause point if BRAND vs TOOL_BRAND mismatches surface — "
    "the worker can't continue without analyst input. We use a threading.Event "
    "as the synchronisation primitive:"
)

code(
    "# Worker side\n"
    "set_state(record, state='mismatch_pending')\n"
    "record.resume_event.wait()                # blocks until set\n"
    "if record.stop_event.is_set():\n"
    "    raise PipelineStopped()\n"
    "\n"
    "# Route side (POST /api/runs/{id}/mismatch/resolve)\n"
    "with record.lock:\n"
    "    record.mismatch_corrections = corrections\n"
    "record.resume_event.set()                 # wakes the worker"
)

p(
    "No polling loop, no busy-waiting. The worker sleeps until the route layer "
    "explicitly wakes it. If the analyst hits 'Cancel run' instead, the route "
    "sets both stop_event and resume_event, and the worker checks stop_event "
    "the moment it wakes."
)

# ── Static export deploy ─────────────────────────────────────────────────────
h2("5. Single-Port Deploy via Static Export")

p(
    "The Next.js frontend is configured for static export "
    "(output: 'export' in next.config.mjs). At build time, npm run build "
    "produces a fully static web/out/ directory — HTML, JS, CSS, fonts, "
    "images — that any static file server can serve."
)

p(
    "FastAPI mounts that directory at the root path so a single uvicorn process "
    "serves both the API and the frontend on one port:"
)

code(
    "WEB_DIST = PROJECT_ROOT / 'web' / 'out'\n"
    "if WEB_DIST.exists():\n"
    "    app.mount('/', StaticFiles(directory=str(WEB_DIST), html=True), name='web')"
)

p(
    "The mount is conditional on web/out existing — in dev, it doesn't (you "
    "run npm run dev separately on :3000) so the static mount is a no-op. In "
    "prod (Docker image), the build stage produces it and the runtime stage "
    "copies it in."
)

# ── NLTK walled-garden bootstrap ─────────────────────────────────────────────
h2("6. The NLTK Walled-Garden Bootstrap")

p(
    "The vendored ml_package calls nltk.download(...) at module import time. "
    "On a sealed network the request fails silently (quiet=True) and "
    "stopwords.words('english') raises LookupError mid-run. The fix has three "
    "pieces:"
)

bullet("Bundle the English corpora in nltk_data/ (~670 KB committed to git)")
bullet("api/_nltk_bootstrap.py prepends nltk_data/ to nltk.data.path and replaces nltk.download with a no-op")
bullet("api/__init__.py imports the bootstrap so it runs before any ml_package import")

code(
    "# api/__init__.py\n"
    "from api import _nltk_bootstrap   # noqa: F401  — side-effect import"
)

p(
    "This pattern lets us fix the issue without touching the vendored code, so "
    "future ml_package refreshes don't blow the fix away."
)

# ── Friendly error mapping ──────────────────────────────────────────────────
h2("7. Friendly Error Mapping")

p(
    "Pipeline failures used to dump tracebacks at the analyst. The new error "
    "module classifies known exception shapes into user-facing dialogs:"
)

code(
    "@dataclass(frozen=True)\n"
    "class FriendlyError:\n"
    "    title:    str\n"
    "    advice:   str\n"
    "    category: str = 'input'\n"
    "\n"
    "def classify(exc: BaseException) -> FriendlyError:\n"
    "    msg = str(exc).lower()\n"
    "    if 'no final sheet' in msg or 'no meta sheet' in msg:\n"
    "        return FriendlyError(\n"
    "            title='Missing required sheet',\n"
    "            advice='The Excel must contain both META and FINAL …',\n"
    "        )\n"
    "    if isinstance(exc, FileNotFoundError):\n"
    "        return FriendlyError(\n"
    "            title='Required project file missing',\n"
    "            advice=f\"The pipeline couldn't find {exc} …\",\n"
    "        )\n"
    "    # … more rules\n"
    "    return FriendlyError(\n"
    "        title='Pipeline failed',\n"
    "        advice='Download the log and share it with support …',\n"
    "        category='server',\n"
    "    )"
)

p(
    "The worker catches every exception, runs classify() on it, and writes the "
    "title + advice + category onto the JobRecord alongside the technical "
    "traceback. JobStatus exposes both — the frontend renders the friendly "
    "version, support escalations have the full trace."
)

# ── Docker ──────────────────────────────────────────────────────────────────
h2("8. Docker — Multi-Stage Build")

p(
    "The Dockerfile is two stages: Node builds the static frontend, Python "
    "serves it. The runtime image only carries the Python dependencies + the "
    "produced web/out — it has no Node, no source TypeScript, no node_modules."
)

code(
    "# Stage 1: Next.js build\n"
    "FROM node:20-alpine AS web-build\n"
    "WORKDIR /web\n"
    "COPY web/package.json web/package-lock.json ./\n"
    "RUN npm ci\n"
    "COPY web/ ./\n"
    "RUN npm run build\n"
    "\n"
    "# Stage 2: Python runtime\n"
    "FROM python:3.11-slim\n"
    "WORKDIR /app\n"
    "COPY requirements.txt ./\n"
    "RUN pip install --no-cache-dir -r requirements.txt\n"
    "COPY api/ ml_package/ phase3_package/ nltk_data/ ./\n"
    "COPY aic_utils.py ./\n"
    "COPY --from=web-build /web/out ./web/out\n"
    "EXPOSE 8000\n"
    "CMD ['uvicorn', 'api.main:app', '--host', '0.0.0.0', '--port', '8000']"
)

p(
    "Build and run:"
)
code(
    "docker build -t aic_v2:latest .\n"
    "docker run -p 8000:8000 aic_v2:latest"
)

callout(
    "The image is ~1.9 GB — most of which is xgboost + scikit-learn + numpy "
    "native binaries. A v2 cleanup could move to python:3.11-slim-bookworm with "
    "more aggressive pruning, but for a tool used by ~20 analysts the size is "
    "fine and the simplicity of one Dockerfile beats a multi-image dance."
)

# ── Concurrent UX ────────────────────────────────────────────────────────────
h2("9. Concurrent-User UX Architecture")

p(
    "The PIPELINE_LOCK means runs serialise. To make that visible instead of "
    "mysterious, the system surfaces three UX surfaces:"
)

h3("Queue position + ETA")

p(
    "The JobRegistry tracks a rolling median of the last 20 successful run "
    "durations per phase. compute_queue_info(record) projects:"
)
bullet("queue_position — count of runs ahead (running + queued earlier)")
bullet("queue_depth — total live runs including yours")
bullet("eta_seconds — median × (position + 1)")

p(
    "These come back on every JobStatus poll while state == 'queued', and the "
    "frontend's ProgressPanel surfaces them as a chip: 'Queued — position 2 of 3 · "
    "~6 min based on recent runs'."
)

h3("Active runs sidebar")

p(
    "GET /api/runs returns a snapshot of every live JobRecord (any state). "
    "The RunsSidebar component polls it every 4 seconds and renders one row "
    "per run with its state chip, phase, and stage label. Click any row to "
    "deep-link to that run's page."
)

h3("Recent runs (localStorage)")

p(
    "When a user starts a run, the run_id + phase + label is recorded into "
    "localStorage. The same RunsSidebar shows those rows under 'Recent on this "
    "device'. If the analyst closes the tab they can come back, see the run, "
    "click it, and resume polling — deep-linkable URLs ?runId=… in both "
    "Phase 1 and Phase 2 pages support this."
)

# ── Testing ──────────────────────────────────────────────────────────────────
h2("10. Testing Strategy")

p("89 / 89 tests pass. Two tiers:")

h3("Fast unit / API tests (~80 tests, < 1 s)")

p(
    "Live in tests/. Heavy ML deps (ml_package, phase3_package) are stubbed in "
    "tests/conftest.py so the FastAPI app can be imported and exercised without "
    "the full ML stack at import time. Covers the registry, QC view shaping, "
    "every API route, validation paths, mismatch enrichment, post-QC, and the "
    "NLTK bootstrap."
)

h3("Integration tests (~5 tests, ~2 s on small fixtures)")

p(
    "Live in tests/integration/. Runs the real ml_package and phase3_package "
    "end-to-end against synthetic fixtures. tests/integration/conftest.py "
    "evicts the fast-test stubs and forces a fresh import of the real packages."
)

p("Stress test (tests/integration/test_stress_e2e.py) covers four scenarios:")
bullet("Phase 1 happy path → finalise → bundle download with verified contents")
bullet("Phase 2 chained off Phase 1 → mismatch resolve → bundle")
bullet("Phase 1 with malformed xlsx → friendly dialog with 'Missing required sheet'")
bullet("Phase 2 with no QC workbook → friendly dialog with title + advice + category")

p("Run the suite:")
code(
    "# Fast suite\n"
    "pytest tests/ --ignore=tests/integration -q\n"
    "\n"
    "# Full suite (needs heavy ML stack)\n"
    "pytest tests/ -q"
)

# ── Higher concepts ─────────────────────────────────────────────────────────
h2("11. Higher Concepts — Why We Built It This Way")

h3("Why FastAPI, not Flask or Django?")

p(
    "FastAPI gives Pydantic validation + OpenAPI docs + async support out of the "
    "box, with less ceremony than Django and more modern primitives than Flask. "
    "The schemas are also our type contract with the frontend (mirrored in "
    "web/lib/types.ts), so changes to JobStatus or QcSheetPayload land in one "
    "Python module and one TypeScript file — symmetrically."
)

h3("Why Next.js static export, not SSR or a SPA framework?")

p(
    "Static export means the frontend is a directory of static files. No Node "
    "process at runtime. No SSR cache invalidation. The deploy artefact is one "
    "Python image that serves both API and frontend. The trade-off is that "
    "anything dynamic (run status) happens in the browser via fetch — that "
    "matches our use case (a polling UI for a long-running job) perfectly."
)

h3("Why one shared PIPELINE_LOCK instead of subprocesses?")

p(
    "Subprocesses are the right long-term answer (each pipeline runs in its "
    "own process, no shared chdir / sys.path). But they add complexity: IPC "
    "for log streaming, pickling for status snapshots, signal handling for "
    "stop. For ~20 analysts on a server with clear queue UX, one shared "
    "process is enough and the code is dramatically simpler. We documented "
    "the upgrade path in the code comments."
)

h3("Why mirror Pydantic schemas in TypeScript by hand?")

p(
    "Codegen tools exist (datamodel-code-generator, openapi-typescript-codegen). "
    "They're a v2 nice-to-have. For 12 type definitions changing once a week "
    "during active development, manual mirroring is faster than the toolchain "
    "setup, the schemas are easy to read side-by-side, and the diff in PRs "
    "shows both files together."
)

page_break()


# ═════════════════════════════════════════════════════════════════════════════
#                              CHAPTER 3
# ═════════════════════════════════════════════════════════════════════════════

h1("Chapter 3 — UI Style Guide and the Case for Having One")

# ── Why a style guide ─────────────────────────────────────────────────────
h2("Why a Style Guide Exists")

p(
    "Three pages (Phase 1, QC wizard, Phase 2 / 3), seven shared components, "
    "and a habit of adding more as the tool grows. Without a style guide, "
    "every page drifts: one card has a 5-pixel padding, another has 6; one "
    "button is brand-purple, another is teal-ish. By the third page the visual "
    "system is incoherent."
)

p(
    "A style guide is a contract: the Circana brand colours, the shadow scale, "
    "the typography, the button system, all live in one place. New pages "
    "compose from the same primitives. Reviewers can spot drift in seconds."
)

# ── Colour system ────────────────────────────────────────────────────────────
h2("The Colour System")

p(
    "Circana brand purple is the spine. The brand-50 to brand-900 scale is "
    "defined in tailwind.config.ts and used everywhere headers, accents, or "
    "interactive elements need brand identity:"
)

code(
    "brand: {\n"
    "  50:  '#F5F0FA',   // lightest tint — background washes\n"
    "  100: '#E8DBF2',\n"
    "  200: '#D2BAE5',\n"
    "  300: '#B594D2',\n"
    "  400: '#8B5DAF',\n"
    "  500: '#6A1A94',\n"
    "  600: '#4E106F',   // primary brand purple\n"
    "  700: '#3A0D54',   // primary CTA, hover\n"
    "  800: '#2E0840',\n"
    "  900: '#1F052B',\n"
    "}"
)

p(
    "Semantic colours sit alongside the brand scale:"
)
bullet("ok / okd — emerald, used for success states and the green primary 'Run' button")
bullet("warn — amber, used for warnings and the mismatch review banner")
bullet("err — red, used for errors and the destructive Stop / Cancel buttons")
bullet("edit — pale lavender, used for edited rows in the QC grid")

callout(
    "The brand scale gives 9 stops so contrast can be tuned without inventing "
    "new colours. Need a hover state for a brand-700 button? Use brand-600. "
    "Need a quiet background for a brand-themed surface? brand-50. Don't "
    "introduce new hex values — extend the scale."
)

# ── Typography ───────────────────────────────────────────────────────────────
h2("Typography")

p(
    "Inter at five weights (400, 500, 600, 700, 800), self-hosted via "
    "next/font/google so the runtime page doesn't depend on Google Fonts CDN. "
    "Inter feature settings cv11, ss01, ss03 are enabled at the html level so "
    "alternates render consistently across browsers. tabular-nums is on by "
    "default so numerics align in tables and progress timers."
)

p(
    "The type scale:"
)
bullet("Display title: text-[36px] sm:text-[42px] font-semibold tracking-tight (page heroes only)")
bullet("Section heading: text-lg font-medium")
bullet("Card title: text-sm font-medium text-zinc-800")
bullet("Body: text-base (17px after the root font-size bump)")
bullet("Metadata / caption: text-xs text-zinc-500")
bullet("Eyebrow: text-[11px] uppercase tracking-[0.14em] text-brand-700/80")

callout(
    "Don't use Tailwind's default text-3xl / text-4xl for hero titles — they're "
    "unscaled, fixed-pixel sizes are deliberate. Use the explicit pixel values "
    "above so the scale stays predictable across the app."
)

# ── Component patterns ───────────────────────────────────────────────────────
h2("Component Patterns")

h3("surface-card")

p(
    "The primary card surface. Translucent white (bg-white/80) with "
    "backdrop-filter saturate(160%) blur(8px), a hairline border, and a "
    "two-stage shadow. Used for ProgressPanel, RunsSidebar, MismatchForm "
    "table, and the Phase 1 / Phase 2 upload sections."
)

code(
    ".surface-card {\n"
    "  @apply rounded-2xl bg-white/80 border;\n"
    "  border-color: var(--hairline);\n"
    "  box-shadow:\n"
    "    0 1px 0 rgba(15, 23, 42, 0.02),\n"
    "    0 6px 24px -12px rgba(15, 23, 42, 0.10);\n"
    "  backdrop-filter: saturate(160%) blur(8px);\n"
    "}"
)

h3("Buttons")

p("Five button variants, all sharing a base size + focus + active behaviour:")
bullet("btn-primary — brand-700 background, white text. Default CTA on QC and Phase 2 download cards.")
bullet("btn-success — ok green, white text. The 'Run pipeline' button on every run-start form.")
bullet("btn-secondary — white background, zinc text, zinc border. Reset / cancel-style alternates.")
bullet("btn-ghost — transparent, zinc text. Tertiary actions.")
bullet("btn-danger-outline — red border, red text, fills red on hover. Stop / cancel buttons.")

p(
    "All buttons share active:scale-[0.98] press feedback, focus-visible "
    "ring-2 + ring-brand-600, and disabled:opacity-50 + cursor-not-allowed. "
    "Add new variants by extending the base — don't fork the colours."
)

h3("Status chips")

p(
    "Pill-shaped, dot + text or just text. Used for queue state, error "
    "categories, run states in the sidebar."
)

code(
    ".chip {\n"
    "  @apply inline-flex items-center gap-1.5 rounded-full px-2.5 py-0.5\n"
    "         text-[11px] font-medium uppercase tracking-wide;\n"
    "}\n"
    ".chip-dot {\n"
    "  @apply h-1.5 w-1.5 rounded-full;\n"
    "}"
)

# ── Layout principles ───────────────────────────────────────────────────────
h2("Layout Principles")

bullet("Max width: max-w-5xl on every page main, max-w-7xl on the QC wizard (the AG Grid is wider).")
bullet("Horizontal padding: px-6 on main, applied at the page level not the section level.")
bullet("Vertical rhythm: pb-12 at the bottom of main, mb-4 / mb-6 between sections.")
bullet("App bar height: h-16 (was h-14 — bumped one notch when the root font-size went from 16 to 17).")
bullet("Page hero: pt-12 pb-7 pseudo-margin so the hero feels generous without crowding the first card.")

# ── Aurora background ───────────────────────────────────────────────────────
h2("Aurora Background and Glass Cards")

p(
    "The page background is two layered radial gradients on a fixed-position "
    "pseudo-element (so it doesn't repaint on scroll). The visual effect is a "
    "soft Circana-purple wash at the top-right + a lavender bleed on the left, "
    "fading into a warm off-white. Cards on top of it are translucent with a "
    "saturate-and-blur backdrop filter — this is what gives them their 'glass' "
    "feel."
)

code(
    "body::before {\n"
    "  content: '';\n"
    "  position: fixed;\n"
    "  inset: 0;\n"
    "  z-index: -1;\n"
    "  background:\n"
    "    radial-gradient(60rem 32rem at 90% -10%,\n"
    "                    rgba(78, 16, 111, 0.16), transparent 60%),\n"
    "    radial-gradient(50rem 28rem at -10% 10%,\n"
    "                    rgba(139, 93, 175, 0.14), transparent 55%),\n"
    "    radial-gradient(40rem 26rem at 50% 110%,\n"
    "                    rgba(232, 219, 242, 0.45), transparent 60%),\n"
    "    linear-gradient(180deg, #FAF7FB 0%, #F6F3F8 100%);\n"
    "}"
)

callout(
    "The aurora is a deliberate choice — flat brand-50 reads as 'corporate "
    "intranet', the gradient reads as 'modern tool'. If you ever want to "
    "dial it back, halve the alphas in the radial-gradient stops."
)

# ── Motion ───────────────────────────────────────────────────────────────────
h2("Motion")

p(
    "Everything respects prefers-reduced-motion. The base motion vocabulary:"
)
bullet("fade-in-up — 360 ms ease entrance, 6 px lift. Applied to the hero and any newly-mounted card.")
bullet("transition-all duration-150 — buttons use this for hover + focus.")
bullet("active:scale-[0.98] — every button presses inward by 2% on click.")

p(
    "Don't add new animations without a clear purpose — page churn from "
    "decorative motion adds friction and accessibility concerns."
)

# ── AG Grid ──────────────────────────────────────────────────────────────────
h2("AG Grid Theme")

p(
    "The QC wizard's grid uses ag-grid-community with the quartz theme "
    "overridden to match Circana brand. Header cells are brand-600 with white "
    "text, alternating row backgrounds use a very pale zinc, and selected "
    "rows use brand-100. Per-row tints come from cellClassRules + Tailwind "
    "classes defined in globals.css (.row-edited, .cell-high-priority, "
    ".cell-low-score-no-ml, .cell-note)."
)

# ── What not to do ──────────────────────────────────────────────────────────
h2("What Not to Do")

bullet("Don't introduce new hex values. Use the brand scale or the semantic colours. If you genuinely need a new colour, add it to tailwind.config.ts so other pages can reuse it.")
bullet("Don't write inline styles for spacing. Use Tailwind utilities so the rhythm stays consistent.")
bullet("Don't rebuild a button from scratch. Use btn-primary / btn-success / etc. — extend the base if you need a new variant.")
bullet("Don't use Tailwind's default font sizes for headings. The hero uses explicit pixel values so the scale is predictable.")
bullet("Don't fetch fonts from Google CDN at runtime. Use next/font/google so they're self-hosted in the static export.")
bullet("Don't forget aria-label on icon-only buttons (the × on QC rows, the Choose button on FileSlot, etc.).")
bullet("Don't ship animations that don't respect prefers-reduced-motion.")

# ── Applying to new pages ───────────────────────────────────────────────────
h2("Applying This Guide to New Applications")

p(
    "Adding a new page (say a settings page) follows the same pattern:"
)
numbered("Wrap content in <Header eyebrow=… title=… subtitle=… /> + <main className='mx-auto max-w-5xl px-6 pb-12'>.")
numbered("Use surface-card for primary content blocks.")
numbered("Use btn-* for every interactive button.")
numbered("Use the brand scale for accents, the semantic palette for state.")
numbered("Use Tailwind utilities for spacing — never inline styles.")
numbered("Add fade-in-up to the main content section so the page enters consistently.")

p(
    "If something doesn't fit the existing primitives, propose an extension to "
    "globals.css or tailwind.config.ts in the same PR — don't accumulate "
    "one-off classes inside the component."
)

# ── Summary ─────────────────────────────────────────────────────────────────
h2("Summary — The Style Decisions at a Glance")

bullet("Circana brand purple (brand-600 = #4E106F) is the primary identity colour")
bullet("Inter, self-hosted via next/font/google, weights 400/500/600/700/800")
bullet("17 px root font-size (bumped from 16 px so everything reads one notch larger)")
bullet("max-w-5xl page widths, 5xl on every page except QC wizard (7xl)")
bullet("Sticky app bar with Circana wordmark + tab nav, then a typographic hero")
bullet("surface-card glass-style primary cards on an aurora gradient background")
bullet("Five-variant button system with consistent focus + press feedback")
bullet("AG Grid with brand-purple headers + per-row tint classes")
bullet("fade-in-up entrance motion, prefers-reduced-motion respected")
bullet("Status chips for queue / state / error category — never plain coloured text")


# ════════════════════════════════════════════════════════════════════════════
# Chapter 4 — Migration: AIC Phase 1 (Streamlit) → AIC v2 (FastAPI + Next.js)
# ════════════════════════════════════════════════════════════════════════════

h1("Chapter 4 — From Streamlit to FastAPI + Next.js")

p(
    "AIC v2 is a ground-up rebuild of the user interface around the same "
    "ml_package and phase3_package code that powered the original AIC Phase 1 "
    "Streamlit app. The pipeline maths is unchanged — every classifier, "
    "lookup, and ensemble call is the same. What changed is everything around "
    "it: how runs are scheduled, how the UI talks to the server, how multiple "
    "analysts share a single deployment, and how the QC grid is rendered. "
    "This chapter exists so an engineer who only knew the Streamlit version "
    "can read the new codebase without surprises."
)

# ── Why ─────────────────────────────────────────────────────────────────────
h2("Why We Refactored")

p(
    "The Streamlit wrapper had hit a ceiling. Three problems compounded:"
)
bullet(
    "Long-running threaded runs collided with Streamlit's rerun model. Every "
    "interaction re-executes the script top-to-bottom, so worker state had to "
    "live in st.session_state and survive reruns. Mid-run UI updates required "
    "ad-hoc hacks (st.empty placeholders, manual reruns)."
)
bullet(
    "The QC grid was constrained by streamlit-aggrid — a thin Streamlit "
    "wrapper around AG Grid Community. We couldn't reach the cellClassRules / "
    "getRowStyle / agSelectCellEditor APIs cleanly, and per-cell tinting for "
    "high-priority and low-score-no-ML rows had to be approximated."
)
bullet(
    "Brand styling lived in load-bearing st.markdown(unsafe_allow_html=True) "
    "blocks. Every page injected raw <style> tags; one typo broke layout. "
    "There was no design-system primitive to reuse, so each page was its own "
    "snowflake."
)

p(
    "All three are symptoms of using a single tool (Streamlit) for jobs it "
    "was never built for: long-running orchestration, custom data grids, and "
    "branded marketing-grade UI. v2 splits those concerns: FastAPI handles "
    "orchestration, Next.js handles UI, and ag-grid-react handles the data "
    "grid — each in its own lane."
)

# ── Architecture ────────────────────────────────────────────────────────────
h2("Architecture: Side-by-Side")

p(
    "The pipeline modules (ml_package, phase3_package, aic_utils) are "
    "vendored as-is. Everything else was rebuilt:"
)

code(
    "AIC Phase 1 (Streamlit)              AIC v2 (FastAPI + Next.js)\n"
    "─────────────────────────────────    ────────────────────────────────────\n"
    "streamlit run streamlit_app.py       uvicorn api.main:app  +  next build\n"
    "Single Python process                FastAPI BFF + static-exported SPA\n"
    "st.session_state for run state       JobRegistry (in-memory + idle TTL)\n"
    "Reruns drive the UI                  React polls /api/runs/{id} every 1s\n"
    "Threaded run, in-process pipeline    Subprocess per pipeline stage\n"
    "_PIPELINE_LOCK (one run at a time)   RUN_SLOTS = BoundedSemaphore(5)\n"
    "streamlit-aggrid                     ag-grid-react (community edition)\n"
    "st.markdown('<style>…')              globals.css + tailwind.config.ts\n"
    "Direct downloads via st.download     /api/runs/{id}/artifacts/{name}\n"
    "Single port (8501)                   Single port (FastAPI mounts web/out)"
)

p(
    "The deploy story is deliberately the same shape — one process, one port "
    "— so the Docker image and operations runbook stay simple. The complexity "
    "lives inside the process, not across services."
)

# ── Concurrency ─────────────────────────────────────────────────────────────
h2("Concurrency: Single Tenant → Five Slots")

p(
    "The biggest behavioural change. The Streamlit app held a "
    "threading.Lock around the whole pipeline call: while one analyst's run "
    "was executing, every other tab sat at 'Waiting for pipeline lock…' until "
    "it finished. Two analysts hitting Run at 9 a.m. meant one of them waited "
    "twenty minutes for the other's job to finish before theirs even started."
)

p(
    "v2 replaces the single lock with a BoundedSemaphore of size 5 and runs "
    "each pipeline stage in a child process via python -m api.run_pipeline. "
    "Five analysts can now run end-to-end simultaneously without contending. "
    "Anyone past the fifth queues at state='queued' with a live ETA projected "
    "from the rolling median of recent runs."
)

callout(
    "Subprocess isolation matters because both ml_package and phase3_package "
    "mutate global state mid-pipeline (os.chdir, sys.path entries). Two "
    "in-process runs would corrupt each other's working directory. One "
    "process per stage side-steps that entirely — the parent just streams "
    "the child's stdout into the JobRecord's log buffer."
)

# ── State & long-running jobs ───────────────────────────────────────────────
h2("State and Long-Running Jobs")

p(
    "Streamlit's session_state is per-tab. Closing the tab loses your run; "
    "opening a second tab to the same URL gives you a fresh empty session "
    "with no idea a job exists. v2 moves the source of truth onto the server:"
)

bullet(
    "JobRegistry holds one JobRecord per run, keyed by a 12-char run_id. "
    "Records carry state (queued, running, qc_ready, mismatch_pending, …), "
    "progress, stage_label, log_lines, and pipeline outputs."
)
bullet(
    "Records are evictable on idle — every API touch refreshes "
    "last_touched, and a background reaper sweeps anything older than the "
    "TTL. An abandoned tab gets its tmpdir cleaned up; an active QC review "
    "stays alive as long as the analyst keeps interacting."
)
bullet(
    "Heavy DataFrames (FINAL, FLAT_FILE_OUT, meta) spill to "
    "phase1_heavy.pkl on disk after Phase 1 completes — the in-memory record "
    "only retains dictEnsemble (the per-attribute frames the QC grid needs). "
    "Five concurrent analysts in QC review hold ~1 GB less RAM than they "
    "would with everything pinned."
)
bullet(
    "Subprocess stdout streams live into record.log_lines, so the UI's log "
    "tail box updates in near-real-time without the parent ever buffering "
    "the whole pipeline output."
)

# ── UI ──────────────────────────────────────────────────────────────────────
h2("UI: From Reruns to Polling")

p(
    "Streamlit's execution model is 'run the script every time anything "
    "changes'. To show progress mid-pipeline, the original code used "
    "st.empty() placeholders and patched them from the worker thread — "
    "fragile, because Streamlit owns the rendering loop. v2 inverts that:"
)

bullet(
    "The Next.js page polls GET /api/runs/{id} every 1 second while a run "
    "is active. The response carries state, progress, stage_label, and the "
    "last 60 log lines. Updates are smooth because the React tree only "
    "re-renders the bits that changed."
)
bullet(
    "Polling stops the moment state hits a terminal (done / error / "
    "stopped) or a paused-on-user-input state (qc_ready / "
    "mismatch_pending). The browser doesn't poll forever — and if you close "
    "the tab, the server has no idea you left, which is fine because "
    "idle-TTL eviction handles cleanup."
)
bullet(
    "Stop is cooperative: POST /api/runs/{id}/stop sets a threading.Event "
    "the worker checks between stages (and which the subprocess receives as "
    "SIGTERM). Same semantic as the Streamlit version, cleaner wiring."
)

# ── QC Grid ────────────────────────────────────────────────────────────────
h2("QC Grid: streamlit-aggrid → ag-grid-react")

p(
    "Both versions use AG Grid Community under the hood, but the integration "
    "shape is completely different:"
)

code(
    "Streamlit                                 v2\n"
    "──────────────────────────────────────    ─────────────────────────────────\n"
    "AgGrid(df, gridOptions=...)               <AgGridReact rowData={…} … />\n"
    "Round-trip every keystroke                Edits stay client-side until save\n"
    "Cell tinting via post-render JS hacks     cellClassRules + Tailwind classes\n"
    "Edits returned as full DataFrame          Diff sent: {row_id, attribute_value}\n"
    "Dropdown options inferred per render      attribute_options shipped once\n"
    "row_flags computed in Streamlit code      Server pre-computes per row"
)

p(
    "The server pre-computes row_flags (priority, low-score-no-ML) and "
    "ships original_values so the React grid can do edit-detection and "
    "cellStyle entirely client-side. Saves are diff-only — only the "
    "(row_id, attribute_value) pairs that actually changed. Finalize "
    "applies the held edits in one pass and writes the xlsx via "
    "ml_package.write_results — same writer the Streamlit version used."
)

# ── Mismatch review ─────────────────────────────────────────────────────────
h2("Mismatch Review: A Pause That Actually Pauses")

p(
    "Phase 2 may surface BRAND vs TOOL_BRAND mismatches that need analyst "
    "judgement. The Streamlit app paused the run by stashing partial state "
    "in session_state and rerunning the script with a different page; the "
    "worker thread effectively died and a fresh one would resume after the "
    "user submitted corrections. v2 keeps the worker alive but releases its "
    "slot:"
)

code(
    "# Worker — Phase A surfaces mismatches\n"
    "with _run_slot(record):                 # slot acquired\n"
    "    interim = run_phase_a(...)          # raises MismatchReviewNeeded\n"
    "# slot released here — others can run while analyst reviews\n"
    "\n"
    "set_state(record, state='mismatch_pending')\n"
    "record.resume_event.wait(timeout=2 * 60 * 60)\n"
    "\n"
    "# Analyst submits — POST /api/runs/{id}/mismatch/resolve\n"
    "with record.lock:\n"
    "    record.mismatch_corrections = corrections\n"
    "record.resume_event.set()\n"
    "\n"
    "with _run_slot(record):                 # fresh slot for Phase B\n"
    "    run_phase_b(interim, corrections, …)"
)

p(
    "Two important properties fall out of this:"
)

bullet(
    "An analyst taking a long lunch during mismatch review doesn't block "
    "anyone else. Their slot has been returned to the pool; the worker is "
    "parked on a threading.Event."
)
bullet(
    "Abandoned reviews self-heal. The 2-hour timeout on resume_event.wait() "
    "means a closed tab eventually marks the run as stopped, the worker "
    "exits, and the JobRecord becomes evictable by idle-TTL. No thread leak."
)

# ── Errors ─────────────────────────────────────────────────────────────────
h2("Error Surfacing: Stack Trace → Title + Advice")

p(
    "Streamlit dumped exceptions raw into st.error — fine for a developer, "
    "intimidating for an analyst staring at a 60-line traceback at 4 p.m. on "
    "a Friday. v2 routes every exception through api.errors.classify, which "
    "returns a (title, advice, category) triple based on the exception type "
    "and message:"
)

code(
    "Original error:\n"
    "  KeyError: 'BRAND_RAW'\n"
    "  File \"ml_package/text_match.py\", line 88, in runTextMatch\n"
    "    df['BRAND_RAW'].fillna('nan')\n"
    "\n"
    "What the analyst sees:\n"
    "  Title:    \"Missing required column\"\n"
    "  Advice:   \"Your input file is missing 'BRAND_RAW'. Re-export from\n"
    "             the source workbook with all four mapping columns.\"\n"
    "  Category: input"
)

p(
    "The full traceback is preserved in the record's error field for "
    "support, but the dialog the analyst sees is a remediation, not a stack "
    "trace. Categories drive the chip colour (input = amber, server = red, "
    "config = blue) and decide whether the error is the user's to fix or "
    "ours."
)

# ── What we kept ────────────────────────────────────────────────────────────
h2("What We Deliberately Kept Untouched")

p(
    "Refactor scope was 'wrapper, not engine'. The following modules ship "
    "byte-identical to the Streamlit version:"
)

bullet("ml_package.mapping_lookup.runLookup")
bullet("ml_package.text_match.runTextMatch")
bullet("ml_package.xgb_classifier.runML")
bullet("ml_package.ensemble.runEnsemble")
bullet("ml_package.write_results.write_results")
bullet("phase3_package.pipeline.run_post_qc and friends")
bullet("aic_utils helper functions")

p(
    "This means a v2 run produces the same xlsx as a Streamlit run for the "
    "same inputs (modulo XGBoost's small OpenMP non-determinism). The "
    "integration test test_qc_edits_round_trip_into_workbook covers this "
    "end-to-end. If a model change is needed, it lands in ml_package and "
    "both deployments inherit it — though only v2 is being maintained going "
    "forward."
)

# ── Onboarding ──────────────────────────────────────────────────────────────
h2("What an Engineer New to v2 Needs to Know")

p(
    "If you knew the Streamlit codebase and are reading v2 for the first "
    "time, the mental shifts are:"
)

numbered(
    "Ignore api/main.py for orientation — it's just route registration. "
    "Start with api/jobs.py (the JobRecord shape and registry), then "
    "api/worker.py (how a run actually executes), then api/run_pipeline.py "
    "(the subprocess entry point that calls into ml_package)."
)
numbered(
    "Routes don't do work. They validate input, create or look up a "
    "JobRecord, and either spawn a worker thread or read a snapshot. All "
    "the business logic is in worker.py and the pipeline modules."
)
numbered(
    "There is no global state outside the JobRegistry singleton. Add new "
    "per-run state by extending JobRecord, not by reaching for a module-"
    "level dict. Add new mutators to api.jobs (set_state, append_log) so "
    "every state change goes through the per-record lock."
)
numbered(
    "The frontend mirrors api/schemas.py manually in web/lib/types.ts — "
    "no codegen. If you change a Pydantic model, update the TypeScript "
    "type in the same PR. Tests will not catch a mismatch."
)
numbered(
    "Tests live in two tiers. tests/ stubs ml_package and runs in <2s — "
    "exercise route logic, state transitions, and edge cases here. "
    "tests/integration/ uses the real ML stack and runs the pipeline "
    "end-to-end — slow but truthful. Both run on every push."
)

callout(
    "If you find yourself adding a workaround that 'feels Streamlit-y' "
    "(global mutable state, mid-function re-renders, st.session_state-"
    "shaped indirection), stop and ask whether the pattern fits FastAPI's "
    "request/response model first. Most of the time the v2-native approach "
    "is half as much code."
)


# ── Save ────────────────────────────────────────────────────────────────────
doc.save(str(OUT_PATH))
print(f"wrote {OUT_PATH} ({OUT_PATH.stat().st_size:,} bytes)")
